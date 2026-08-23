from io import BytesIO

from docx import Document


def generate_resume_docx(profile):
    document = Document()
    document.add_heading(profile.full_name or "Resume", 0)
    document.add_paragraph(profile.headline or "")

    if profile.about:
        document.add_heading('О себе', level=1)
        document.add_paragraph(profile.about)

    if profile.institution:
        document.add_heading('Образование', level=1)
        document.add_paragraph(f"{profile.institution} - {profile.specialty}")

    f = BytesIO()
    document.save(f)
    return f.getvalue()
